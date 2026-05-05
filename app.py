import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
import subprocess
import os
import sys
import threading
import shutil
import uuid as _uuid
import tempfile
import fitz  # PyMuPDF
from PIL import Image as _PILImage
from PIL import ImageFilter as _PILImageFilter
from PIL import ImageOps as _PILImageOps

try:
    from pdf2docx import Converter
    HAS_PDF2DOCX = True
except ImportError:
    HAS_PDF2DOCX = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ── 字體（macOS 系統字體）────────────────────────────────────────────────────
FONT = "SF Pro Text"

# ── Color palette ────────────────────────────────────────────────────────────
PRIMARY           = "#0058bc"
PRIMARY_CONTAINER = "#0070eb"
ON_PRIMARY        = "#ffffff"
SURFACE           = "#f9f9fe"
SURFACE_LOWEST    = "#ffffff"
SURFACE_LOW       = "#f3f3f8"
SURFACE_CONTAINER = "#ededf2"
SURFACE_HIGH      = "#e8e8ed"
SURFACE_HIGHEST   = "#e2e2e7"
ON_SURFACE        = "#1a1c1f"
ON_SURFACE_VAR    = "#414755"
OUTLINE_VAR       = "#c1c6d7"
SIDEBAR_BG        = "#dde1ea"
GREEN_TAG         = "#72fe88"
GREEN_TAG_TEXT    = "#002107"
SUCCESS_COLOR     = "#006b27"
ERROR_COLOR       = "#ba1a1a"

PAGE_NAMES = {
    "tasks":    "工作",
    "history":  "歷史記錄",
    "settings": "設定",
}


# ── 全域任務列表 ──────────────────────────────────────────────────────────────
TASKS: list[dict] = []

def _task_add(name: str) -> str:
    tid = _uuid.uuid4().hex[:8]
    TASKS.append({"id": tid, "name": name, "status": "running", "message": ""})
    return tid

def _task_update(tid: str, *, status: str = None, message: str = None):
    for t in TASKS:
        if t["id"] == tid:
            if status is not None:
                t["status"] = status
            if message is not None:
                t["message"] = message
            break


# ── 工具函式 ─────────────────────────────────────────────────────────────────
def get_gs_path():
    # 優先使用 PyInstaller 打包進 .app 的 Ghostscript binary
    if hasattr(sys, "_MEIPASS"):
        bundled = os.path.join(sys._MEIPASS, "gs")
        if os.path.exists(bundled):
            return bundled
    # 開發環境 fallback：依序嘗試常見系統路徑
    for candidate in ["/opt/homebrew/bin/gs", "/usr/local/bin/gs"]:
        if os.path.exists(candidate):
            return candidate
    return shutil.which("gs")


def get_tesseract_path():
    if hasattr(sys, "_MEIPASS"):
        bundled = os.path.join(sys._MEIPASS, "tesseract")
        if os.path.exists(bundled):
            return bundled
    for candidate in ["/opt/homebrew/bin/tesseract", "/usr/local/bin/tesseract"]:
        if os.path.exists(candidate):
            return candidate
    return shutil.which("tesseract")


def get_tessdata_dir():
    if hasattr(sys, "_MEIPASS"):
        bundled = os.path.join(sys._MEIPASS, "tessdata")
        if os.path.isdir(bundled):
            return bundled
    for candidate in ["/opt/homebrew/share/tessdata", "/usr/local/share/tessdata"]:
        if os.path.isdir(candidate):
            return candidate
    return None


def get_pdf_text_stats(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        page_count = doc.page_count
        chars_by_page = [len(page.get_text("text").strip()) for page in doc]
        doc.close()
        total_chars = sum(chars_by_page)
        pages_with_text = sum(1 for count in chars_by_page if count >= 20)
        return page_count, total_chars, pages_with_text
    except Exception:
        return 0, 0, 0


def should_use_ocr(pdf_path):
    page_count, total_chars, pages_with_text = get_pdf_text_stats(pdf_path)
    if page_count == 0:
        return False
    avg_chars = total_chars / page_count
    text_page_ratio = pages_with_text / page_count
    return total_chars < 80 or avg_chars < 30 or text_page_ratio < 0.3


def validate_ocr_ready():
    if not HAS_DOCX:
        return "缺少 python-docx 套件，無法產生 OCR Word 檔。"

    tesseract = get_tesseract_path()
    if not tesseract:
        return "找不到 Tesseract OCR，無法處理掃描型 PDF。"

    tessdata = get_tessdata_dir()
    if not tessdata:
        return "找不到 Tesseract 語言資料目錄。"

    missing = []
    for lang in ["chi_tra", "eng"]:
        if not os.path.exists(os.path.join(tessdata, f"{lang}.traineddata")):
            missing.append(lang)
    if missing:
        return "缺少 OCR 語言包：" + "、".join(missing) + "。請安裝繁體中文與英文語言包。"
    return None


def preprocess_image_for_ocr(input_path, output_path):
    image = _PILImage.open(input_path).convert("L")
    image = _PILImageOps.autocontrast(image)
    image = image.filter(_PILImageFilter.SHARPEN)
    image = image.point(lambda pixel: 255 if pixel > 185 else 0)
    image.save(output_path, dpi=(300, 300))


def clean_ocr_lines(text):
    lines = []
    previous_blank = False
    for raw_line in text.replace("\t", " ").splitlines():
        line = " ".join(raw_line.strip().split())
        if not line:
            if not previous_blank and lines:
                lines.append("")
            previous_blank = True
            continue
        lines.append(line)
        previous_blank = False
    while lines and not lines[-1]:
        lines.pop()
    return lines


def ocr_pdf_to_docx(pdf_path, output_path, status_callback=None):
    tesseract = get_tesseract_path()
    tessdata = get_tessdata_dir()
    document = Document()
    source_name = os.path.basename(pdf_path)
    document.add_heading(os.path.splitext(source_name)[0], level=1)

    with fitz.open(pdf_path) as pdf, tempfile.TemporaryDirectory() as tmpdir:
        total = pdf.page_count
        for index, page in enumerate(pdf, start=1):
            if status_callback:
                status_callback(f"正在 OCR 辨識第 {index}/{total} 頁，請稍候…")

            zoom = 300 / 72
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
            image_path = os.path.join(tmpdir, f"page-{index}.png")
            processed_path = os.path.join(tmpdir, f"page-{index}-ocr.png")
            pix.save(image_path)
            preprocess_image_for_ocr(image_path, processed_path)

            cmd = [
                tesseract,
                processed_path,
                "stdout",
                "-l",
                "chi_tra+eng",
                "--tessdata-dir",
                tessdata,
                "--oem",
                "1",
                "--psm",
                "4",
                "-c",
                "preserve_interword_spaces=1",
            ]
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=240,
            )
            if result.returncode != 0:
                raise RuntimeError(result.stderr.strip() or "OCR 辨識失敗")

            if index > 1:
                document.add_page_break()
            document.add_heading(f"第 {index} 頁", level=2)

            lines = clean_ocr_lines(result.stdout)
            if lines:
                for line in lines:
                    if line:
                        document.add_paragraph(line)
                    else:
                        document.add_paragraph("")
            else:
                document.add_paragraph("（此頁未辨識到文字）")

    document.save(output_path)


def format_size(b):
    if b < 1024:
        return f"{b} B"
    if b < 1024 * 1024:
        return f"{b / 1024:.1f} KB"
    return f"{b / (1024 * 1024):.1f} MB"


def analyze_pdf(pdf_path):
    """回傳 (頁數, 含圖片, 圖片佔比)"""
    try:
        doc = fitz.open(pdf_path)
        page_count = doc.page_count
        file_size = os.path.getsize(pdf_path)
        total_image_bytes = 0
        seen = set()
        for page in doc:
            for img in page.get_images():
                xref = img[0]
                if xref not in seen:
                    seen.add(xref)
                    try:
                        data = doc.extract_image(xref)
                        total_image_bytes += len(data.get("image", b""))
                    except Exception:
                        pass
        doc.close()
        image_ratio = min(total_image_bytes / max(file_size, 1), 0.95)
        return page_count, len(seen) > 0, image_ratio
    except Exception:
        return 0, False, 0.3


def estimate_size(file_size, image_ratio, quality):
    """回傳 (預估大小 bytes, 縮減比例)"""
    img_reduction = {"ebook": 0.45, "balanced": 0.65, "screen": 0.80}[quality]
    reduction = image_ratio * img_reduction + (1 - image_ratio) * 0.08
    reduction = min(max(reduction, 0.0), 0.95)
    return int(file_size * (1 - reduction)), reduction



# ── 壓縮對話框 ───────────────────────────────────────────────────────────────
class CompressDialog(ctk.CTkToplevel):
    def __init__(self, parent, input_path=None):
        super().__init__(parent)
        self.title("縮小 PDF 容量")
        self.geometry("520x560")
        self.resizable(False, False)
        self.configure(fg_color=SURFACE)
        self.grab_set()
        self.lift()
        self.focus_force()

        self.input_path = input_path
        self._task_id: str | None = None
        self.quality_var = tk.StringVar(value="ebook")
        self._est_labels: dict[str, ctk.CTkLabel] = {}
        self._build()
        if input_path:
            self._refresh_file_label()
            threading.Thread(target=self._analyze, daemon=True).start()

    def _build(self):
        # 標題
        ctk.CTkLabel(self, text="縮小 PDF 容量",
                     font=ctk.CTkFont(FONT, 20, "bold"),
                     text_color=ON_SURFACE).pack(anchor="w", padx=28, pady=(28, 2))
        ctk.CTkLabel(self, text="選擇品質預設並選取 PDF 檔案。",
                     font=ctk.CTkFont(FONT, 13),
                     text_color=ON_SURFACE_VAR).pack(anchor="w", padx=28)

        # 選擇檔案
        file_frame = ctk.CTkFrame(self, fg_color=SURFACE_HIGH, corner_radius=12)
        file_frame.pack(fill="x", padx=28, pady=14)
        self.file_label = ctk.CTkLabel(file_frame, text="尚未選擇檔案",
                                        font=ctk.CTkFont(FONT, 13),
                                        text_color=ON_SURFACE_VAR)
        self.file_label.pack(side="left", padx=16, pady=12)
        ctk.CTkButton(file_frame, text="瀏覽", width=72,
                      fg_color=PRIMARY, hover_color=PRIMARY_CONTAINER,
                      text_color=ON_PRIMARY, corner_radius=8,
                      font=ctk.CTkFont(FONT, 13),
                      command=self._browse).pack(side="right", padx=10, pady=8)

        # PDF 資訊面板
        self.info_frame = ctk.CTkFrame(self, fg_color=SURFACE_LOW, corner_radius=10)
        self.info_frame.pack(fill="x", padx=28, pady=(0, 8))
        self.info_label = ctk.CTkLabel(
            self.info_frame,
            text="選擇 PDF 後顯示檔案資訊",
            font=ctk.CTkFont(FONT, 12),
            text_color=ON_SURFACE_VAR)
        self.info_label.pack(padx=16, pady=10, anchor="w")

        # 品質設定
        ctk.CTkLabel(self, text="品質設定",
                     font=ctk.CTkFont(FONT, 14, "bold"),
                     text_color=ON_SURFACE).pack(anchor="w", padx=28, pady=(6, 4))

        presets = [
            ("高品質（150 DPI） — 縮減 20–50%", "ebook"),
            ("平衡（96 DPI）　  — 縮減 40–70%", "balanced"),
            ("最小化（72 DPI）  — 縮減 60–90%", "screen"),
        ]
        for label, val in presets:
            row = ctk.CTkFrame(self, fg_color="transparent")
            row.pack(fill="x", padx=28, pady=2)
            row.grid_columnconfigure(0, weight=1)

            ctk.CTkRadioButton(row, text=label,
                               variable=self.quality_var, value=val,
                               font=ctk.CTkFont(FONT, 13),
                               fg_color=PRIMARY,
                               hover_color=PRIMARY_CONTAINER).grid(
                row=0, column=0, sticky="w")

            est = ctk.CTkLabel(row, text="",
                               font=ctk.CTkFont(FONT, 12),
                               text_color=ON_SURFACE_VAR)
            est.grid(row=0, column=1, sticky="e")
            self._est_labels[val] = est

        # 狀態
        self.status_label = ctk.CTkLabel(self, text="",
                                          font=ctk.CTkFont(FONT, 13),
                                          text_color=ON_SURFACE_VAR,
                                          wraplength=460, justify="left")
        self.status_label.pack(anchor="w", padx=28, pady=(12, 4))

        # 執行按鈕
        self.action_btn = ctk.CTkButton(
            self, text="開始壓縮  →",
            fg_color=PRIMARY, hover_color=PRIMARY_CONTAINER,
            text_color=ON_PRIMARY,
            font=ctk.CTkFont(FONT, 14, "bold"),
            height=46, corner_radius=12,
            command=self._start)
        self.action_btn.pack(fill="x", padx=28, pady=(0, 28))

    def _browse(self):
        path = filedialog.askopenfilename(
            title="選擇 PDF 檔案", filetypes=[("PDF 檔案", "*.pdf")])
        if path:
            self.input_path = path
            self._refresh_file_label()
            threading.Thread(target=self._analyze, daemon=True).start()

    def _refresh_file_label(self):
        name = os.path.basename(self.input_path)
        size = os.path.getsize(self.input_path)
        self.file_label.configure(text=f"{name}  （{format_size(size)}）")
        self.info_label.configure(text="分析中…")
        for lbl in self._est_labels.values():
            lbl.configure(text="計算中…")

    def _analyze(self):
        pages, has_images, img_ratio = analyze_pdf(self.input_path)
        file_size = os.path.getsize(self.input_path)
        self.after(0, lambda: self._update_info(pages, has_images, img_ratio, file_size))

    def _update_info(self, pages, has_images, img_ratio, file_size):
        img_text = "含圖片" if has_images else "純文字"
        page_text = f"{pages} 頁" if pages else "—"
        self.info_label.configure(
            text=f"共 {page_text}  ·  {img_text}  ·  {format_size(file_size)}")

        for quality, lbl in self._est_labels.items():
            est_bytes, reduction = estimate_size(file_size, img_ratio, quality)
            lbl.configure(text=f"→ 約 {format_size(est_bytes)}（-{reduction*100:.0f}%）")

    def _start(self):
        if not self.input_path:
            messagebox.showwarning("未選擇檔案", "請先選擇一個 PDF 檔案。", parent=self)
            return
        if not get_gs_path():
            messagebox.showerror("找不到 Ghostscript",
                                 "請先安裝 Ghostscript：\n\nbrew install ghostscript",
                                 parent=self)
            return
        self.action_btn.configure(state="disabled", text="壓縮中…")
        self.status_label.configure(text="正在壓縮，請稍候…", text_color=ON_SURFACE_VAR)
        self._task_id = _task_add(f"{os.path.basename(self.input_path)} — PDF 壓縮")
        threading.Thread(target=self._compress, daemon=True).start()

    def _compress(self):
        quality   = self.quality_var.get()
        base, ext = os.path.splitext(self.input_path)
        output    = f"{base}_compressed{ext}"
        gs        = get_gs_path()

        if quality == "balanced":
            cmd = [gs, "-sDEVICE=pdfwrite", "-dCompatibilityLevel=1.4",
                   "-dNOPAUSE", "-dQUIET", "-dBATCH",
                   "-dDownsampleColorImages=true", "-dColorImageResolution=96",
                   "-dDownsampleGrayImages=true",  "-dGrayImageResolution=96",
                   "-dColorImageDownsampleThreshold=1.0",
                   "-dGrayImageDownsampleThreshold=1.0",
                   f"-sOutputFile={output}", self.input_path]
        else:
            setting = "/ebook" if quality == "ebook" else "/screen"
            cmd = [gs, "-sDEVICE=pdfwrite", "-dCompatibilityLevel=1.4",
                   f"-dPDFSETTINGS={setting}",
                   "-dNOPAUSE", "-dQUIET", "-dBATCH",
                   f"-sOutputFile={output}", self.input_path]
        try:
            r = subprocess.run(cmd, capture_output=True, text=True, timeout=180)
            if r.returncode == 0 and os.path.exists(output):
                orig = os.path.getsize(self.input_path)
                new  = os.path.getsize(output)
                pct  = (1 - new / orig) * 100
                self.after(0, lambda: self._done(output, orig, new, pct))
            else:
                self.after(0, lambda: self._error(r.stderr or "未知錯誤"))
        except Exception as e:
            self.after(0, lambda: self._error(str(e)))

    def _done(self, output, orig, new, pct):
        self.action_btn.configure(state="normal", text="開始壓縮  →")
        msg = (f"✓ 完成！\n"
               f"原始：{format_size(orig)}  →  壓縮後：{format_size(new)}  （-{pct:.0f}%）\n"
               f"輸出：{os.path.basename(output)}")
        self.status_label.configure(text=msg, text_color=SUCCESS_COLOR)
        if self._task_id:
            _task_update(self._task_id, status="done", message=msg)

    def _error(self, msg):
        self.action_btn.configure(state="normal", text="開始壓縮  →")
        self.status_label.configure(text=f"錯誤：{msg[:120]}", text_color=ERROR_COLOR)
        if self._task_id:
            _task_update(self._task_id, status="error", message=f"錯誤：{msg[:120]}")


# ── 轉換對話框 ───────────────────────────────────────────────────────────────
class ConvertDialog(ctk.CTkToplevel):
    def __init__(self, parent, input_path=None):
        super().__init__(parent)
        self.title("轉換為 Word")
        self.geometry("500x300")
        self.resizable(False, False)
        self.configure(fg_color=SURFACE)
        self.grab_set()
        self.lift()
        self.focus_force()

        self.input_path = input_path
        self._task_id: str | None = None
        self._build()
        if input_path:
            self._refresh_file_label()

    def _build(self):
        ctk.CTkLabel(self, text="轉換為 Word",
                     font=ctk.CTkFont(FONT, 20, "bold"),
                     text_color=ON_SURFACE).pack(anchor="w", padx=28, pady=(28, 2))
        ctk.CTkLabel(self, text="將 PDF 轉換為可編輯的 Word 文件格式。",
                     font=ctk.CTkFont(FONT, 13),
                     text_color=ON_SURFACE_VAR).pack(anchor="w", padx=28)

        file_frame = ctk.CTkFrame(self, fg_color=SURFACE_HIGH, corner_radius=12)
        file_frame.pack(fill="x", padx=28, pady=16)
        self.file_label = ctk.CTkLabel(file_frame, text="尚未選擇檔案",
                                        font=ctk.CTkFont(FONT, 13),
                                        text_color=ON_SURFACE_VAR)
        self.file_label.pack(side="left", padx=16, pady=12)
        ctk.CTkButton(file_frame, text="瀏覽", width=72,
                      fg_color=SURFACE_HIGHEST, text_color=ON_SURFACE,
                      hover_color=OUTLINE_VAR, corner_radius=8,
                      font=ctk.CTkFont(FONT, 13),
                      command=self._browse).pack(side="right", padx=10, pady=8)

        self.status_label = ctk.CTkLabel(self, text="",
                                          font=ctk.CTkFont(FONT, 13),
                                          text_color=ON_SURFACE_VAR,
                                          wraplength=440, justify="left")
        self.status_label.pack(anchor="w", padx=28, pady=(4, 8))

        self.action_btn = ctk.CTkButton(
            self, text="開始轉換  →",
            fg_color=SURFACE_HIGHEST, text_color=ON_SURFACE,
            hover_color=OUTLINE_VAR,
            font=ctk.CTkFont(FONT, 14, "bold"),
            height=46, corner_radius=12,
            command=self._start)
        self.action_btn.pack(fill="x", padx=28, pady=(0, 28))

    def _browse(self):
        path = filedialog.askopenfilename(
            title="選擇 PDF 檔案", filetypes=[("PDF 檔案", "*.pdf")])
        if path:
            self.input_path = path
            self._refresh_file_label()

    def _refresh_file_label(self):
        name = os.path.basename(self.input_path)
        size = os.path.getsize(self.input_path)
        self.file_label.configure(text=f"{name}  （{format_size(size)}）")

    def _start(self):
        if not self.input_path:
            messagebox.showwarning("未選擇檔案", "請先選擇一個 PDF 檔案。", parent=self)
            return
        self.action_btn.configure(state="disabled", text="轉換中…")
        self.status_label.configure(text="正在分析 PDF…", text_color=ON_SURFACE_VAR)
        self._task_id = _task_add(f"{os.path.basename(self.input_path)} — 轉換為 Word")
        threading.Thread(target=self._convert, daemon=True).start()

    def _convert(self):
        try:
            output = os.path.splitext(self.input_path)[0] + ".docx"
            if should_use_ocr(self.input_path):
                err = validate_ocr_ready()
                if err:
                    raise RuntimeError(err)
                self._set_status("偵測為掃描 PDF，使用 OCR 辨識文字…\nOCR 可能需要較長時間，請稍候…")
                ocr_pdf_to_docx(self.input_path, output, self._set_status)
                self.after(0, lambda: self._done(output, "已使用 OCR 辨識"))
            else:
                if not HAS_PDF2DOCX:
                    raise RuntimeError("缺少 pdf2docx 套件，無法使用一般轉換。")
                self._set_status("偵測到可編輯文字，使用一般轉換…")
                cv = Converter(self.input_path)
                try:
                    cv.convert(output)
                finally:
                    cv.close()
                self.after(0, lambda: self._done(output, "已使用一般轉換"))
        except Exception as e:
            self.after(0, lambda: self._error(str(e)))

    def _set_status(self, msg):
        self.after(0, lambda: self.status_label.configure(
            text=msg, text_color=ON_SURFACE_VAR))
        if self._task_id:
            _task_update(self._task_id, status="running", message=msg)

    def _done(self, output, mode):
        self.action_btn.configure(state="normal", text="開始轉換  →")
        msg = f"✓ 完成！{mode}\n輸出：{os.path.basename(output)}"
        self.status_label.configure(
            text=msg,
            text_color=SUCCESS_COLOR)
        if self._task_id:
            _task_update(self._task_id, status="done", message=msg)

    def _error(self, msg):
        self.action_btn.configure(state="normal", text="開始轉換  →")
        self.status_label.configure(text=f"錯誤：{msg[:120]}", text_color=ERROR_COLOR)
        if self._task_id:
            _task_update(self._task_id, status="error", message=f"錯誤：{msg[:120]}")


# ── 側邊欄導覽按鈕 ────────────────────────────────────────────────────────────
class NavButton(ctk.CTkButton):
    def __init__(self, master, label, icon, active=False, **kw):
        super().__init__(
            master,
            text=f"  {icon}   {label}",
            anchor="w",
            height=40,
            corner_radius=10,
            fg_color=SURFACE_LOWEST if active else "transparent",
            text_color=PRIMARY if active else ON_SURFACE_VAR,
            hover_color=SURFACE_LOW,
            font=ctk.CTkFont(FONT, 14, "bold" if active else "normal"),
            **kw
        )

    def set_active(self, active: bool):
        self.configure(
            fg_color=SURFACE_LOWEST if active else "transparent",
            text_color=PRIMARY if active else ON_SURFACE_VAR,
            font=ctk.CTkFont(FONT, 14, "bold" if active else "normal"),
        )


# ── 主程式 ────────────────────────────────────────────────────────────────────
class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")

        self.title("PDF 輕巧工具箱")
        self.geometry("960x660")
        self.minsize(820, 560)
        self.configure(fg_color=SURFACE)

        self._nav_buttons: dict[str, NavButton] = {}
        self._active_page = "home"

        self._build_sidebar()
        self._build_main()
        self._show_home()

    # ── 側邊欄 ────────────────────────────────────────────────────────────────
    def _build_sidebar(self):
        self.sidebar = ctk.CTkFrame(self, width=224, corner_radius=0,
                                     fg_color=SIDEBAR_BG)
        self.sidebar.pack(side="left", fill="y")
        self.sidebar.pack_propagate(False)

        logo = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        logo.pack(fill="x", padx=18, pady=(32, 24))
        ctk.CTkLabel(logo, text="PDF 輕巧工具箱",
                     font=ctk.CTkFont(FONT, 17, "bold"),
                     text_color=ON_SURFACE).pack(anchor="w")
        ctk.CTkLabel(logo, text="PDF 工作區",
                     font=ctk.CTkFont(FONT, 11),
                     text_color=ON_SURFACE_VAR).pack(anchor="w")

        nav = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        nav.pack(fill="x", padx=12)

        items = [
            ("home",     "首頁",    "⌂"),
            ("tasks",    "工作",    "✓"),
            ("history",  "歷史記錄", "◷"),
            ("settings", "設定",    "⚙"),
        ]
        for key, label, icon in items:
            btn = NavButton(nav, label, icon,
                            active=(key == "home"),
                            command=lambda k=key: self._navigate(k))
            btn.pack(fill="x", pady=2)
            self._nav_buttons[key] = btn

        profile = ctk.CTkFrame(self.sidebar, fg_color=SURFACE_LOW, corner_radius=12)
        profile.pack(side="bottom", fill="x", padx=12, pady=16)
        ctk.CTkLabel(profile, text="Patrick",
                     font=ctk.CTkFont(FONT, 13, "bold"),
                     text_color=ON_SURFACE).pack(anchor="w", padx=14, pady=(10, 0))
        ctk.CTkLabel(profile, text="本機版本",
                     font=ctk.CTkFont(FONT, 11),
                     text_color=ON_SURFACE_VAR).pack(anchor="w", padx=14, pady=(0, 10))

    # ── 主內容區 ──────────────────────────────────────────────────────────────
    def _build_main(self):
        self.canvas = ctk.CTkScrollableFrame(self, fg_color=SURFACE, corner_radius=0)
        self.canvas.pack(side="left", fill="both", expand=True)

    def _clear_canvas(self):
        for w in self.canvas.winfo_children():
            w.destroy()

    # ── 導覽切換 ──────────────────────────────────────────────────────────────
    def _navigate(self, page: str):
        self._nav_buttons[self._active_page].set_active(False)
        self._active_page = page
        self._nav_buttons[page].set_active(True)
        self._clear_canvas()
        if page == "home":
            self._show_home()
        elif page == "tasks":
            self._show_tasks()
        else:
            self._show_placeholder(PAGE_NAMES.get(page, page))

    # ── 首頁 ──────────────────────────────────────────────────────────────────
    def _show_home(self):
        hdr = ctk.CTkFrame(self.canvas, fg_color="transparent")
        hdr.pack(fill="x", padx=52, pady=(52, 36))
        ctk.CTkLabel(hdr, text="PDF 輕巧工具箱",
                     font=ctk.CTkFont(FONT, 44, "bold"),
                     text_color=ON_SURFACE).pack(anchor="w")
        ctk.CTkLabel(hdr,
                     text="智慧型文件工作區，精準壓縮並轉換您的 PDF 檔案。",
                     font=ctk.CTkFont(FONT, 14),
                     text_color=ON_SURFACE_VAR,
                     wraplength=520).pack(anchor="w", pady=(4, 0))

        row = ctk.CTkFrame(self.canvas, fg_color="transparent")
        row.pack(fill="x", padx=52, pady=(0, 28))
        row.grid_columnconfigure((0, 1), weight=1, uniform="card")

        self._card_compress(row)
        self._card_word(row)
        self._drop_zone()

    def _card_compress(self, parent):
        card = ctk.CTkFrame(parent, fg_color=SURFACE_LOWEST, corner_radius=16,
                             border_width=1, border_color=OUTLINE_VAR)
        card.grid(row=0, column=0, padx=(0, 14), sticky="nsew", ipady=4)

        inner = ctk.CTkFrame(card, fg_color="transparent")
        inner.pack(fill="both", expand=True, padx=28, pady=28)

        ctk.CTkLabel(inner, text="縮小 PDF 容量",
                     font=ctk.CTkFont(FONT, 20, "bold"),
                     text_color=ON_SURFACE).pack(anchor="w", pady=(16, 6))
        ctk.CTkLabel(inner,
                     text="智慧優化文件品質，大幅縮小檔案容量，方便傳送與分享。",
                     font=ctk.CTkFont(FONT, 13),
                     text_color=ON_SURFACE_VAR,
                     wraplength=220, justify="left").pack(anchor="w", fill="x")

        bottom = ctk.CTkFrame(inner, fg_color="transparent")
        bottom.pack(fill="x", pady=(24, 0))
        bottom.grid_columnconfigure(0, weight=1)
        ctk.CTkLabel(bottom, text="高效壓縮",
                     font=ctk.CTkFont(FONT, 11, "bold"),
                     fg_color=GREEN_TAG, text_color=GREEN_TAG_TEXT,
                     corner_radius=99, padx=10, pady=4).grid(row=0, column=0, sticky="w")
        ctk.CTkButton(bottom, text="壓縮 →",
                      fg_color=PRIMARY, hover_color=PRIMARY_CONTAINER,
                      text_color=ON_PRIMARY,
                      font=ctk.CTkFont(FONT, 13, "bold"),
                      width=100, height=38, corner_radius=10,
                      command=lambda: CompressDialog(self)).grid(row=0, column=1, sticky="e")

    def _card_word(self, parent):
        card = ctk.CTkFrame(parent, fg_color=SURFACE_LOW, corner_radius=16,
                             border_width=1, border_color=OUTLINE_VAR)
        card.grid(row=0, column=1, padx=(14, 0), sticky="nsew", ipady=4)

        inner = ctk.CTkFrame(card, fg_color="transparent")
        inner.pack(fill="both", expand=True, padx=28, pady=28)

        ctk.CTkLabel(inner, text="轉換為 Word",
                     font=ctk.CTkFont(FONT, 20, "bold"),
                     text_color=ON_SURFACE).pack(anchor="w", pady=(16, 6))
        ctk.CTkLabel(inner,
                     text="高精度轉換，完整保留 PDF 的版面、字體與表格。",
                     font=ctk.CTkFont(FONT, 13),
                     text_color=ON_SURFACE_VAR,
                     wraplength=220, justify="left").pack(anchor="w", fill="x")

        bottom = ctk.CTkFrame(inner, fg_color="transparent")
        bottom.pack(fill="x", pady=(24, 0))
        bottom.grid_columnconfigure(0, weight=1)
        ctk.CTkLabel(bottom, text="版面完整",
                     font=ctk.CTkFont(FONT, 11, "bold"),
                     fg_color=SURFACE_CONTAINER, text_color=ON_SURFACE_VAR,
                     corner_radius=99, padx=10, pady=4).grid(row=0, column=0, sticky="w")
        ctk.CTkButton(bottom, text="轉換 →",
                      fg_color=SURFACE_HIGHEST, text_color=ON_SURFACE,
                      hover_color=OUTLINE_VAR,
                      font=ctk.CTkFont(FONT, 13, "bold"),
                      width=100, height=38, corner_radius=10,
                      command=lambda: ConvertDialog(self)).grid(row=0, column=1, sticky="e")

    def _drop_zone(self):
        zone = ctk.CTkFrame(self.canvas, fg_color=SURFACE_LOW, corner_radius=16,
                             border_width=2, border_color=OUTLINE_VAR)
        zone.pack(fill="x", padx=52, pady=(0, 52))

        inner = ctk.CTkFrame(zone, fg_color="transparent")
        inner.pack(pady=44)

        ctk.CTkLabel(inner, text="⬆", font=ctk.CTkFont(size=52),
                     text_color="#9db8e8").pack()
        ctk.CTkLabel(inner, text="點擊選擇 PDF 檔案",
                     font=ctk.CTkFont(FONT, 16, "bold"),
                     text_color=ON_SURFACE_VAR).pack(pady=(10, 4))
        ctk.CTkLabel(inner, text="或將 PDF 拖曳至 Dock 上的應用程式圖示",
                     font=ctk.CTkFont(FONT, 13),
                     text_color=ON_SURFACE_VAR).pack()

        tags = ctk.CTkFrame(inner, fg_color="transparent")
        tags.pack(pady=20)
        for t in ["🔒  本機處理", "⚡  快速安全"]:
            ctk.CTkLabel(tags, text=t,
                         font=ctk.CTkFont(FONT, 11, "bold"),
                         fg_color=SURFACE_HIGH, text_color=ON_SURFACE_VAR,
                         corner_radius=99, padx=14, pady=6).pack(side="left", padx=6)

        # 點擊開啟檔案選擇器
        def on_click(_event=None):
            path = filedialog.askopenfilename(
                title="選擇 PDF 檔案", filetypes=[("PDF 檔案", "*.pdf")])
            if path:
                CompressDialog(self, path)

        for w in [zone, inner] + inner.winfo_children() + tags.winfo_children():
            w.bind("<Button-1>", on_click)
            w.configure(cursor="hand2")

        # macOS 原生：從 Finder 拖曳至 Dock 圖示
        try:
            self.createcommand("::tk::mac::OpenDocument", self._on_open_document)
        except Exception:
            pass

    def _on_open_document(self, *args):
        """macOS 拖曳至 Dock 圖示或 Finder 開啟"""
        for path in args:
            if isinstance(path, str) and path.lower().endswith(".pdf") and os.path.isfile(path):
                CompressDialog(self, path)
                return

    # ── 工作頁面 ──────────────────────────────────────────────────────────────
    def _show_tasks(self):
        ctk.CTkLabel(self.canvas, text="工作",
                     font=ctk.CTkFont(FONT, 36, "bold"),
                     text_color=ON_SURFACE).pack(anchor="w", padx=52, pady=(60, 4))
        ctk.CTkLabel(self.canvas, text="正在進行與已完成的工作。",
                     font=ctk.CTkFont(FONT, 14),
                     text_color=ON_SURFACE_VAR).pack(anchor="w", padx=52, pady=(0, 28))

        if not TASKS:
            ctk.CTkLabel(self.canvas, text="目前沒有進行中的工作。",
                         font=ctk.CTkFont(FONT, 14),
                         text_color=ON_SURFACE_VAR).pack(anchor="w", padx=52)
            return

        for t in reversed(TASKS):
            card = ctk.CTkFrame(self.canvas, fg_color=SURFACE_LOWEST,
                                corner_radius=14,
                                border_width=1, border_color=OUTLINE_VAR)
            card.pack(fill="x", padx=52, pady=(0, 14))

            inner = ctk.CTkFrame(card, fg_color="transparent")
            inner.pack(fill="x", padx=24, pady=18)

            # 標題列
            if t["status"] == "running":
                icon = "⏳"
                icon_color = PRIMARY
            elif t["status"] == "done":
                icon = "✓"
                icon_color = SUCCESS_COLOR
            else:
                icon = "✗"
                icon_color = ERROR_COLOR

            title_row = ctk.CTkFrame(inner, fg_color="transparent")
            title_row.pack(fill="x")
            ctk.CTkLabel(title_row, text=icon,
                         font=ctk.CTkFont(FONT, 14, "bold"),
                         text_color=icon_color).pack(side="left")
            ctk.CTkLabel(title_row, text=f"  {t['name']}",
                         font=ctk.CTkFont(FONT, 14, "bold"),
                         text_color=ON_SURFACE).pack(side="left")

            # 進度條（進行中顯示不確定式）
            if t["status"] == "running":
                pb = ctk.CTkProgressBar(inner,
                                        mode="indeterminate",
                                        fg_color=SURFACE_HIGH,
                                        progress_color=PRIMARY,
                                        height=8,
                                        corner_radius=4)
                pb.pack(fill="x", pady=(12, 4))
                pb.start()
            else:
                pb = ctk.CTkProgressBar(inner,
                                        mode="determinate",
                                        fg_color=SURFACE_HIGH,
                                        progress_color=SUCCESS_COLOR if t["status"] == "done" else ERROR_COLOR,
                                        height=8,
                                        corner_radius=4)
                pb.pack(fill="x", pady=(12, 4))
                pb.set(1.0)

            # 訊息
            if t["message"]:
                msg_color = SUCCESS_COLOR if t["status"] == "done" else (ERROR_COLOR if t["status"] == "error" else ON_SURFACE_VAR)
                ctk.CTkLabel(inner, text=t["message"],
                             font=ctk.CTkFont(FONT, 12),
                             text_color=msg_color,
                             wraplength=700, justify="left").pack(anchor="w", pady=(4, 0))

        # 若有執行中的任務，每秒重新渲染
        if any(t["status"] == "running" for t in TASKS):
            self.canvas.after(1000, self._refresh_tasks_if_active)

    def _refresh_tasks_if_active(self):
        if self._active_page == "tasks":
            self._clear_canvas()
            self._show_tasks()

    # ── 預留頁面 ──────────────────────────────────────────────────────────────
    def _show_placeholder(self, name):
        ctk.CTkLabel(self.canvas, text=name,
                     font=ctk.CTkFont(FONT, 36, "bold"),
                     text_color=ON_SURFACE).pack(anchor="w", padx=52, pady=(60, 8))
        ctk.CTkLabel(self.canvas, text="此功能即將推出。",
                     font=ctk.CTkFont(FONT, 14),
                     text_color=ON_SURFACE_VAR).pack(anchor="w", padx=52)


if __name__ == "__main__":
    app = App()
    app.mainloop()
